import json
import io
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from datetime import datetime
from docx import Document

from ...db.database import get_db
from ...db.models import User, Conversation, Draft
from ...core.auth import get_current_user
from ...schemas.schemas import QuestionRequest, DocumentRequest
from ...services.vector_service import vector_service
from ...services.bns_service import bns_service

router = APIRouter(tags=["Citizen - RAG & Drafting"])

import re

def html_to_plain_text(html_content: str) -> str:
    if not html_content:
        return ""
    # Convert common paragraph/line break tags to newlines
    text = html_content.replace("<br>", "\n").replace("<br/>", "\n").replace("<br />", "\n")
    text = re.sub(r'</div>', '\n', text)
    text = re.sub(r'</p>', '\n', text)
    # Strip all other HTML tags
    text = re.sub(r'<[^>]*>', '', text)
    # Decode common HTML entities
    text = text.replace("&nbsp;", " ").replace("&lt;", "<").replace("&gt;", ">").replace("&amp;", "&")
    return text

def text_to_docx(text: str, title: str):
    doc = Document()
    doc.add_heading(title, 0)
    clean_text = html_to_plain_text(text)
    for line in clean_text.split('\n'):
        doc.add_paragraph(line)
    
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

@router.get("/download-docx/{draft_id}")
def download_docx(draft_id: int, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    draft = db.query(Draft).filter(Draft.id == draft_id, Draft.user_id == current_user.id).first()
    if not draft: raise HTTPException(status_code=404, detail="Not found")
    
    file_stream = text_to_docx(draft.content, draft.title)
    return StreamingResponse(
        file_stream, 
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={draft.title.replace(' ', '_')}.docx"}
    )

@router.post("/ask")
def ask_question(body: QuestionRequest, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if not body.question.strip():
        raise HTTPException(status_code=400, detail="Empty question")
    
    result = vector_service.get_citizen_answer(
        body.question, 
        "general", 
        "Provide a helpful legal answer.",
        language=body.language or "English"
    )
    
    db.add(Conversation(user_id=current_user.id, question=body.question, answer=result["answer"]))
    db.commit()
    return result

@router.post("/generate-document")
def generate_doc(body: DocumentRequest, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    desc = body.fields.get("description", "")
    content = vector_service.generate_legal_document(body.doc_type, desc)
    
    draft = Draft(
        user_id=current_user.id,
        doc_type=body.doc_type,
        title=f"{body.doc_type.replace('_', ' ').title()} - {datetime.now().strftime('%d %b %Y')}",
        content=content,
        form_data=json.dumps(body.fields)
    )
    db.add(draft)
    db.commit()
    return {"content": content, "title": draft.title, "draft_id": draft.id}

@router.get("/history")
def get_history(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    return db.query(Conversation).filter(Conversation.user_id == current_user.id).all()

@router.delete("/history/{item_id}")
def delete_history(item_id: int, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    item = db.query(Conversation).filter(Conversation.id == item_id, Conversation.user_id == current_user.id).first()
    if not item: raise HTTPException(status_code=404, detail="Not found")
    db.delete(item)
    db.commit()
    return {"message": "Deleted"}

@router.get("/drafts")
def get_drafts(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    return db.query(Draft).filter(Draft.user_id == current_user.id).all()

@router.get("/drafts/{draft_id}")
def get_draft(draft_id: int, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    draft = db.query(Draft).filter(Draft.id == draft_id, Draft.user_id == current_user.id).first()
    if not draft: raise HTTPException(status_code=404, detail="Not found")
    return draft

@router.delete("/drafts/{draft_id}")
def delete_draft(draft_id: int, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    draft = db.query(Draft).filter(Draft.id == draft_id, Draft.user_id == current_user.id).first()
    if not draft: raise HTTPException(status_code=404, detail="Not found")
    db.delete(draft)
    db.commit()
    return {"message": "Deleted"}

@router.get("/bns-helper")
def bns_lookup(query: str, current_user: User = Depends(get_current_user)):
    if not query.strip():
        raise HTTPException(status_code=400, detail="Empty query")
    return bns_service.lookup(query)

@router.put("/drafts/{draft_id}")
def update_draft(draft_id: int, body: dict, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    draft = db.query(Draft).filter(Draft.id == draft_id, Draft.user_id == current_user.id).first()
    if not draft: 
        raise HTTPException(status_code=404, detail="Not found")
    
    content = body.get("content")
    if content is None:
        raise HTTPException(status_code=400, detail="Missing content")
    
    draft.content = content
    draft.updated_at = datetime.now()
    db.commit()
    return {"message": "Updated successfully", "content": draft.content}
